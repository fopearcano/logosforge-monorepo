"""Export project data to JSON, Markdown, CSV, DOCX, Fountain, PDF, HTML, or FDX."""

import csv
import io
import json
import xml.etree.ElementTree as ET

from logosforge.db import Database


def _build_timeline_section(db: Database, project_id: int, scenes: list) -> dict:
    """Timeline lanes + event links. Event positions/colours are already carried
    on each scene (order_index/plotline/color_label); links reference scenes by
    their 1-based export order so they survive re-import."""
    order_by_sid = {s.id: i + 1 for i, s in enumerate(scenes)}
    title_by_sid = {s.id: s.title for s in scenes}
    lanes = [
        {
            "name": ln.name,
            "color_label": ln.color_label,
            "order_index": ln.order_index,
            "collapsed": ln.collapsed,
        }
        for ln in db.get_timeline_lanes(project_id)
    ]
    links = [
        {
            "source_order": order_by_sid.get(ln.source_scene_id),
            "target_order": order_by_sid.get(ln.target_scene_id),
            "source_title": title_by_sid.get(ln.source_scene_id, ""),
            "target_title": title_by_sid.get(ln.target_scene_id, ""),
            "color_label": ln.color_label,
            "link_type": ln.link_type,
            "label": ln.label,
        }
        for ln in db.get_timeline_links(project_id)
    ]
    structure_links = [
        {
            "source_order": order_by_sid.get(sl.source_scene_id),
            "source_title": title_by_sid.get(sl.source_scene_id, ""),
            "target_type": sl.target_type,
            "target_ref": sl.target_ref,
        }
        for sl in db.get_all_timeline_structure_links(project_id)
    ]
    return {"lanes": lanes, "links": links,
            "structure_links": structure_links}


def _gather_project_data(db: Database, project_id: int) -> dict:
    project = db.get_project_by_id(project_id)

    characters = db.get_all_characters(project_id)
    places = db.get_all_places(project_id)
    notes = db.get_all_notes(project_id)
    scenes = db.get_all_scenes(project_id)

    char_name_by_id = {c.id: c.name for c in characters}
    place_name_by_id = {p.id: p.name for p in places}

    scene_list = []
    for i, scene in enumerate(scenes):
        char_ids = db.get_scene_character_ids(scene.id)
        place_ids = db.get_scene_place_ids(scene.id)
        char_states = db.get_scene_character_states(scene.id)

        scene_list.append(
            {
                "title": scene.title,
                "summary": scene.summary,
                "synopsis": scene.synopsis,
                "goal": scene.goal,
                "conflict": scene.conflict,
                "outcome": scene.outcome,
                "content": scene.content,
                "beat": scene.beat,
                "tags": [t.strip() for t in scene.tags.split(",") if t.strip()] if scene.tags else [],
                "order_index": i + 1,
                "act": scene.act,
                "chapter": scene.chapter,
                "plotline": scene.plotline,
                "color_label": scene.color_label,
                # -- Screenplay-engine fields (empty for non-screenplay scenes)
                "slugline": scene.slugline,
                "location": scene.location,
                "interior_exterior": scene.interior_exterior,
                "time_of_day": scene.time_of_day,
                "estimated_duration_minutes": scene.estimated_duration_minutes,
                "visual_objective": scene.visual_objective,
                "dramatic_turn": scene.dramatic_turn,
                "blocking_notes": scene.blocking_notes,
                "subtext_notes": scene.subtext_notes,
                "setup_payoff_links": scene.setup_payoff_links,
                "montage_group": scene.montage_group,
                "cinematic_pacing": scene.cinematic_pacing,
                "continuity_notes": scene.continuity_notes,
                # -- Screenplay PSYKE extensions ------------------------
                "visible_conflict": scene.visible_conflict,
                "hidden_conflict": scene.hidden_conflict,
                "emotional_turn": scene.emotional_turn,
                "who_knows_what": scene.who_knows_what,
                "physical_action": scene.physical_action,
                "visual_symbolism": scene.visual_symbolism,
                "characters": [
                    char_name_by_id[cid]
                    for cid in char_ids
                    if cid in char_name_by_id
                ],
                "places": [
                    place_name_by_id[pid]
                    for pid in place_ids
                    if pid in place_name_by_id
                ],
                "character_states": [
                    {"character": char_name_by_id[cid], "state": state}
                    for cid, state in char_states
                    if cid in char_name_by_id
                ],
            }
        )

    psyke_entries = db.get_all_psyke_entries(project_id)
    psyke_name_by_id = {e.id: e.name for e in psyke_entries}
    psyke_names_lower = {e.name.lower() for e in psyke_entries}
    characters = [c for c in characters if c.name.lower() not in psyke_names_lower]
    places = [p for p in places if p.name.lower() not in psyke_names_lower]

    scene_title_by_id = {s.id: s.title for s in scenes}

    psyke_list = []
    for e in psyke_entries:
        typed_related = db.get_typed_related_psyke_entries(e.id)
        progressions = db.get_psyke_progressions(e.id)
        psyke_list.append({
            "name": e.name,
            "entry_type": e.entry_type,
            "aliases": e.aliases,
            "notes": e.notes,
            "is_global": e.is_global,
            "details": db.get_psyke_entry_details(e.id),
            "related_entries": [r.name for r, _ in typed_related],
            "typed_relations": [
                {"name": r.name, "relation_type": rtype}
                for r, rtype in typed_related
            ],
            "progressions": [
                {
                    "text": p.text,
                    "scene_title": scene_title_by_id.get(p.scene_id, "")
                    if p.scene_id
                    else "",
                    "sort_order": p.sort_order,
                }
                for p in progressions
            ],
        })

    continuity_items = []
    for scene in scenes:
        for item in db.get_continuity_for_scene(scene.id):
            continuity_items.append({
                "scene_title": scene.title,
                "memory_type": item.memory_type,
                "target": item.target,
                "value": item.value,
            })

    outline_nodes = db.get_outline_nodes(project_id)
    children_map: dict[int | None, list] = {}
    for node in outline_nodes:
        children_map.setdefault(node.parent_id, []).append(node)

    def _build_outline_tree(parent_id: int | None) -> list[dict]:
        children = children_map.get(parent_id, [])
        children.sort(key=lambda n: (n.sort_order, n.id))
        return [
            {
                "title": n.title,
                "description": n.description,
                "children": _build_outline_tree(n.id),
            }
            for n in children
        ]

    from logosforge.quantum_outliner.persistence import export_quantum_state

    from logosforge.project_compat import (
        get_project_narrative_engine,
        get_project_writing_format,
    )
    from logosforge.writing_modes import get_project_writing_mode
    data = {
        "project": {
            "title": project.title if project else "Untitled",
            "description": project.description if project else "",
            "format_mode": (project.format_mode if project else "novel") or "novel",
            "writing_mode": get_project_writing_mode(project),
            "narrative_engine": get_project_narrative_engine(project),
            "default_writing_format": get_project_writing_format(project),
        },
        "characters": [
            {"name": c.name, "description": c.description} for c in characters
        ],
        "places": [
            {"name": p.name, "description": p.description} for p in places
        ],
        "notes": [
            {
                "title": n.title,
                "content": n.content,
                "tags": n.tags,
                "pinned": n.pinned,
                "psyke_links": [
                    psyke_name_by_id[eid]
                    for eid in db.get_note_psyke_links(n.id)
                    if eid in psyke_name_by_id
                ],
                "scene_links": [
                    scene_title_by_id[sid]
                    for sid in db.get_note_scene_links(n.id)
                    if sid in scene_title_by_id
                ],
                "structure_links": [
                    {"type": ttype, "ref": ref}
                    for ttype, ref in db.get_note_structure_links(n.id)
                ],
            }
            for n in notes
        ],
        "scenes": scene_list,
        "chapters": [
            {
                "act": ch.act,
                "title": ch.title,
                "summary": ch.summary,
                "content": ch.content,
                "order_index": ch.order_index,
            }
            for ch in db.get_chapters(project_id)
        ],
        "psyke_entries": psyke_list,
        "outline": _build_outline_tree(None),
        "continuity": continuity_items,
        "plot_timeline": _build_timeline_section(db, project_id, scenes),
    }

    quantum = export_quantum_state(db, project_id)
    if quantum is not None:
        data["quantum_state"] = quantum

    return data


def export_json(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)
    return json.dumps(data, indent=2, ensure_ascii=False)


def export_screenplay_diagnostics_json(db: Database, project_id: int) -> str:
    """Phase 10C — per-scene deterministic screenplay diagnostics as JSON.

    Additive and read-only; does not touch the existing export paths. Includes
    the project's writing mode and a per-scene report list.
    """
    from logosforge.screenplay_diagnostics import analyze_project
    from logosforge.writing_modes import get_project_writing_mode

    project = db.get_project_by_id(project_id)
    payload = {
        "project": {
            "title": project.title if project else "Untitled",
            "writing_mode": get_project_writing_mode(project),
        },
        "scenes": [r.to_dict() for r in analyze_project(db, project_id)],
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def export_setup_payoff_report_json(db: Database, project_id: int) -> str:
    """Phase 10D — project setup/payoff candidate report as JSON (read-only)."""
    from logosforge.screenplay_setup_payoff import analyze_setup_payoff
    from logosforge.writing_modes import get_project_writing_mode

    project = db.get_project_by_id(project_id)
    payload = {
        "project": {
            "title": project.title if project else "Untitled",
            "writing_mode": get_project_writing_mode(project),
        },
        "setup_payoff": analyze_setup_payoff(db, project_id).to_dict(),
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def export_subtext_report_json(db: Database, project_id: int) -> str:
    """Phase 10D — per-scene subtext report as JSON (read-only)."""
    from logosforge.screenplay_subtext import analyze_subtext_project
    from logosforge.writing_modes import get_project_writing_mode

    project = db.get_project_by_id(project_id)
    payload = {
        "project": {
            "title": project.title if project else "Untitled",
            "writing_mode": get_project_writing_mode(project),
        },
        "scenes": [r.to_dict() for r in analyze_subtext_project(db, project_id)],
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def export_screenplay_graph_json(db: Database, project_id: int) -> str:
    """Phase 10E — screenplay story-link graph as JSON (read-only)."""
    from logosforge.screenplay_graph import build_screenplay_graph
    from logosforge.writing_modes import get_project_writing_mode

    project = db.get_project_by_id(project_id)
    graph = build_screenplay_graph(db, project_id)
    payload = graph.to_dict()
    payload["project"] = {
        "title": project.title if project else "Untitled",
        "writing_mode": get_project_writing_mode(project),
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def export_story_links_json(db: Database, project_id: int) -> str:
    """Phase 10E — persisted confirmed/tracked story links as JSON (read-only)."""
    from logosforge.writing_modes import get_project_writing_mode

    project = db.get_project_by_id(project_id)
    try:
        links = db.get_story_links(project_id)
    except Exception:
        links = []
    payload = {
        "schema_version": 1,
        "project": {
            "title": project.title if project else "Untitled",
            "writing_mode": get_project_writing_mode(project),
        },
        "story_links": [
            {
                "id": l.id, "link_type": l.link_type, "label": l.label,
                "status": l.status, "source_type": l.source_type,
                "source_id": l.source_id, "source_scene_id": l.source_scene_id,
                "target_type": l.target_type, "target_id": l.target_id,
                "target_scene_id": l.target_scene_id, "evidence": l.evidence,
                "confidence": l.confidence,
            }
            for l in links
        ],
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def export_markdown(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)
    lines: list[str] = []

    # Project header
    lines.append(f"# {data['project']['title']}")
    from logosforge.writing_modes import mode_label
    lines.append("")
    lines.append(f"*Writing Mode: {mode_label(data['project'].get('writing_mode'))}*")
    if data["project"]["description"]:
        lines.append("")
        lines.append(data["project"]["description"])

    # Characters
    lines.append("")
    lines.append("## Characters")
    if data["characters"]:
        for char in data["characters"]:
            lines.append("")
            lines.append(f"### {char['name']}")
            if char["description"]:
                lines.append("")
                lines.append(char["description"])
    else:
        lines.append("")
        lines.append("No characters.")

    # Places
    lines.append("")
    lines.append("## Places")
    if data["places"]:
        for place in data["places"]:
            lines.append("")
            lines.append(f"### {place['name']}")
            if place["description"]:
                lines.append("")
                lines.append(place["description"])
    else:
        lines.append("")
        lines.append("No places.")

    # Notes
    lines.append("")
    lines.append("## Notes")
    if data["notes"]:
        for note in data["notes"]:
            lines.append("")
            lines.append(f"### {note['title']}")
            if note["content"]:
                lines.append("")
                lines.append(note["content"])
    else:
        lines.append("")
        lines.append("No notes.")

    # Scenes
    lines.append("")
    lines.append("## Scenes")
    if data["scenes"]:
        for scene in data["scenes"]:
            lines.append("")
            lines.append(f"### {scene['order_index']}. {scene['title']}")
            if scene["act"]:
                lines.append(f"- **Act:** {scene['act']}")
            if scene["chapter"]:
                lines.append(f"- **Chapter:** {scene['chapter']}")
            if scene["plotline"]:
                lines.append(f"- **Plotline:** {scene['plotline']}")
            if scene["beat"]:
                lines.append(f"- **Beat:** {scene['beat']}")
            if scene["tags"]:
                lines.append(f"- **Tags:** {', '.join(scene['tags'])}")
            if scene["characters"]:
                lines.append(f"- **Characters:** {', '.join(scene['characters'])}")
            if scene["places"]:
                lines.append(f"- **Places:** {', '.join(scene['places'])}")
            if scene["summary"]:
                lines.append("")
                lines.append(scene["summary"])
            if scene["synopsis"]:
                lines.append("")
                lines.append(f"**Synopsis:** {scene['synopsis']}")
            if scene["goal"]:
                lines.append(f"- **Goal:** {scene['goal']}")
            if scene["conflict"]:
                lines.append(f"- **Conflict:** {scene['conflict']}")
            if scene["outcome"]:
                lines.append(f"- **Outcome:** {scene['outcome']}")
            if scene["character_states"]:
                lines.append("")
                lines.append("**Character States:**")
                for cs in scene["character_states"]:
                    lines.append(f"- {cs['character']}: {cs['state']}")
            if scene["content"]:
                lines.append("")
                lines.append("---")
                lines.append("")
                lines.append(scene["content"])
    else:
        lines.append("")
        lines.append("No scenes.")

    lines.append("")
    return "\n".join(lines)


def export_outline_markdown(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)
    lines: list[str] = []

    lines.append(f"# {data['project']['title']}")

    scenes = data["scenes"]
    if not scenes:
        lines.append("")
        lines.append("No scenes to display.")
        lines.append("")
        return "\n".join(lines)

    chapter_groups = _group_scenes_by_chapter(scenes)

    for chapter_name, group_scenes in chapter_groups:
        lines.append("")
        if chapter_name:
            lines.append(f"## {chapter_name}")
        else:
            lines.append("## Uncategorized")

        for scene in group_scenes:
            lines.append("")
            lines.append(f"### {scene['order_index']}. {scene['title']}")
            if scene["act"]:
                lines.append(f"- **Act:** {scene['act']}")
            if scene["plotline"]:
                lines.append(f"- **Plotline:** {scene['plotline']}")
            if scene["beat"]:
                lines.append(f"- **Beat:** {scene['beat']}")
            if scene["tags"]:
                lines.append(f"- **Tags:** {', '.join(scene['tags'])}")
            if scene["characters"]:
                lines.append(f"- **Characters:** {', '.join(scene['characters'])}")
            if scene["places"]:
                lines.append(f"- **Places:** {', '.join(scene['places'])}")
            if scene["summary"]:
                lines.append("")
                lines.append(scene["summary"])
            if scene["synopsis"]:
                lines.append("")
                lines.append(f"**Synopsis:** {scene['synopsis']}")
            if scene["goal"]:
                lines.append(f"- **Goal:** {scene['goal']}")
            if scene["conflict"]:
                lines.append(f"- **Conflict:** {scene['conflict']}")
            if scene["outcome"]:
                lines.append(f"- **Outcome:** {scene['outcome']}")
            if scene["character_states"]:
                lines.append("")
                lines.append("**Character States:**")
                for cs in scene["character_states"]:
                    lines.append(f"- {cs['character']}: {cs['state']}")

    lines.append("")
    return "\n".join(lines)


def export_csv_scenes(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(
        ["order_index", "title", "summary", "synopsis", "goal", "conflict", "outcome",
         "beat", "tags", "act", "chapter", "plotline", "characters", "places"]
    )
    for scene in data["scenes"]:
        writer.writerow(
            [
                scene["order_index"],
                scene["title"],
                scene["summary"],
                scene["synopsis"],
                scene["goal"],
                scene["conflict"],
                scene["outcome"],
                scene["beat"],
                ", ".join(scene["tags"]),
                scene["act"],
                scene["chapter"],
                scene["plotline"],
                ", ".join(scene["characters"]),
                ", ".join(scene["places"]),
            ]
        )
    return output.getvalue()


def _group_scenes_by_chapter(scenes: list[dict]) -> list[tuple[str, list[dict]]]:
    groups: list[tuple[str, list[dict]]] = []
    current_chapter: str | None = None
    current_group: list[dict] = []

    for scene in scenes:
        chapter = scene["chapter"] if scene["chapter"] else ""
        if chapter != current_chapter:
            if current_group:
                groups.append((current_chapter or "", current_group))
            current_chapter = chapter
            current_group = [scene]
        else:
            current_group.append(scene)
    if current_group:
        groups.append((current_chapter or "", current_group))

    return groups


def _scene_body(scene: dict) -> str:
    return scene["content"] or scene["synopsis"] or scene["summary"] or ""


def _scene_starts_with_heading(scene: dict) -> bool:
    """True if the scene content already opens with a parsed scene heading.

    Used to avoid emitting a duplicate metadata slug line on export.
    """
    raw = _scene_body(scene)
    if not raw.strip():
        return False
    try:
        from logosforge.screenplay_blocks import parse_screenplay_text
        blocks = parse_screenplay_text(raw)
        return bool(blocks) and blocks[0].element_type == "scene_heading"
    except Exception:
        return False


def _screenplay_body(scene: dict, *, fountain: bool = False,
                     include_notes: bool = True) -> str:
    """Render a scene body as classified screenplay text (Phase 10B/10F).

    Parses the flat scene content into screenplay blocks and serializes them so
    character cues / transitions / scene headings are uppercased and
    parentheticals normalized. Text-preserving; falls back to the raw body if
    parsing yields nothing. When *include_notes* is False, ``note`` blocks are
    dropped (production export).
    """
    raw = _scene_body(scene)
    if not raw.strip():
        return ""
    try:
        from logosforge.screenplay_blocks import (
            parse_screenplay_text,
            serialize_blocks,
            to_fountain,
        )
        blocks = parse_screenplay_text(raw)
        if not blocks:
            return raw
        if not include_notes:
            blocks = [b for b in blocks if b.element_type != "note"]
        return to_fountain(blocks) if fountain else serialize_blocks(blocks)
    except Exception:
        return raw


def export_screenplay(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)
    return _format_text(data, "screenplay")


def export_manuscript(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)
    return _format_text(data, "novel")


def export_formatted_text(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)
    return _format_text(data, _get_fmt(data))


def _is_script_format(fmt: str) -> bool:
    return fmt in ("screenplay", "series", "stage_script", "graphic_novel")


def _format_text(data: dict, fmt: str) -> str:
    if fmt in ("screenplay", "series"):
        return _fmt_screenplay_text(data, fmt)
    if fmt == "stage_script":
        return _fmt_stage_script_text(data)
    if fmt == "graphic_novel":
        return _fmt_graphic_novel_text(data)
    return _fmt_novel_text(data)


def _slug_line(scene: dict) -> str:
    if scene["places"]:
        place_str = ", ".join(scene["places"]).upper()
        return f"INT. {place_str} — {scene['title'].upper()}"
    return scene["title"].upper()


def _fmt_novel_text(data: dict) -> str:
    lines: list[str] = []
    title = data["project"]["title"]
    lines.append(title)
    lines.append("=" * len(title))
    lines.append("")

    if not data["scenes"]:
        lines.append("No scenes.")
        return "\n".join(lines)

    chapter_groups = _group_scenes_by_chapter(data["scenes"])

    for chapter_name, group_scenes in chapter_groups:
        if chapter_name:
            lines.append("")
            lines.append(chapter_name)
            lines.append("-" * len(chapter_name))
            lines.append("")

        for scene in group_scenes:
            lines.append(scene["title"])
            lines.append("")
            body = _scene_body(scene)
            if body:
                lines.append(body)
            lines.append("")
            lines.append("")

    return "\n".join(lines)


def _fmt_screenplay_text(data: dict, fmt: str) -> str:
    lines: list[str] = []
    title = data["project"]["title"]
    lines.append(title.upper())
    lines.append("=" * len(title))
    # Phase 10A — record the canonical writing mode in the screenplay header.
    from logosforge.writing_modes import mode_label
    lines.append(f"Writing Mode: {mode_label(data['project'].get('writing_mode'))}")
    lines.append("")
    lines.append("")

    current_act = None
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if fmt == "series" and act and act != current_act:
            current_act = act
            lines.append("")
            lines.append(f"        {act.upper()}")
            lines.append("")

        if not _scene_starts_with_heading(scene):
            lines.append(_slug_line(scene))
            lines.append("")
        body = _screenplay_body(scene)
        if body:
            lines.append(body)
            lines.append("")
        lines.append("")

    return "\n".join(lines)


def _fmt_stage_script_text(data: dict) -> str:
    lines: list[str] = []
    title = data["project"]["title"]
    lines.append(f"        {title.upper()}")
    lines.append("")
    lines.append("")

    current_act = None
    current_scene_num = 0
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if act and act != current_act:
            current_act = act
            current_scene_num = 0
            lines.append("")
            lines.append(f"        {act.upper()}")
            lines.append("")

        current_scene_num += 1
        lines.append(f"        SCENE {current_scene_num}")
        lines.append("")

        body = _scene_body(scene)
        if body:
            lines.append(body)
            lines.append("")
        lines.append("")

    return "\n".join(lines)


def _fmt_graphic_novel_text(data: dict) -> str:
    lines: list[str] = []
    title = data["project"]["title"]
    lines.append(title.upper())
    lines.append("")

    page_num = 0
    for scene in data["scenes"]:
        page_num += 1
        lines.append(f"PAGE {page_num}")
        lines.append("")
        body = _scene_body(scene)
        if body:
            lines.append(body)
            lines.append("")
        lines.append("")

    return "\n".join(lines)


# -- DOCX export (format-aware) -----------------------------------------------

_SCRIPT_FONT_DOCX = "Courier New"
_PROSE_FONT_DOCX = "Times New Roman"


def _get_fmt(data: dict) -> str:
    return data["project"].get("format_mode", "novel")


def export_docx_manuscript(db: Database, project_id: int, path: str) -> None:
    data = _gather_project_data(db, project_id)
    fmt = _get_fmt(data)
    if fmt in ("screenplay", "series"):
        _docx_screenplay(data, path, fmt)
    elif fmt == "stage_script":
        _docx_stage_script(data, path)
    elif fmt == "graphic_novel":
        _docx_graphic_novel(data, path)
    else:
        _docx_novel(data, path)


def _docx_title_page(doc, title: str, font_name: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.name = font_name
    title_para.paragraph_format.space_after = Pt(0)
    doc.add_page_break()


def _docx_novel(data: dict, path: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _PROSE_FONT_DOCX
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    _docx_title_page(doc, data["project"]["title"], _PROSE_FONT_DOCX)

    if not data["scenes"]:
        doc.add_paragraph("No scenes.")
        doc.save(path)
        return

    chapter_groups = _group_scenes_by_chapter(data["scenes"])
    chapter_num = 0
    for chapter_name, group_scenes in chapter_groups:
        if chapter_name:
            chapter_num += 1
            heading = f"Chapter {chapter_num}: {chapter_name}"
        else:
            heading = "Scenes"
        ch_para = doc.add_paragraph()
        ch_para.paragraph_format.space_before = Pt(24)
        ch_run = ch_para.add_run(heading)
        ch_run.bold = True
        ch_run.font.size = Pt(16)
        ch_run.font.name = _PROSE_FONT_DOCX

        for scene in group_scenes:
            scene_para = doc.add_paragraph()
            scene_para.paragraph_format.space_before = Pt(12)
            scene_run = scene_para.add_run(scene["title"])
            scene_run.italic = True
            scene_run.font.size = Pt(12)
            scene_run.font.name = _PROSE_FONT_DOCX

            body = _scene_body(scene)
            if body:
                _add_content_paragraphs(doc, body, _PROSE_FONT_DOCX)

    doc.save(path)


def _docx_screenplay(data: dict, path: str, fmt: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _SCRIPT_FONT_DOCX
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    _docx_title_page(doc, data["project"]["title"], _SCRIPT_FONT_DOCX)

    if not data["scenes"]:
        doc.add_paragraph("No scenes.")
        doc.save(path)
        return

    current_act = None
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if fmt == "series" and act and act != current_act:
            current_act = act
            act_para = doc.add_paragraph()
            act_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            act_para.paragraph_format.space_before = Pt(24)
            act_run = act_para.add_run(act.upper())
            act_run.bold = True
            act_run.font.name = _SCRIPT_FONT_DOCX

        slug_para = doc.add_paragraph()
        slug_para.paragraph_format.space_before = Pt(24)
        slug_para.paragraph_format.space_after = Pt(12)
        slug_run = slug_para.add_run(_slug_line(scene))
        slug_run.bold = True
        slug_run.font.name = _SCRIPT_FONT_DOCX

        body = _scene_body(scene)
        if body:
            _add_content_paragraphs(doc, body, _SCRIPT_FONT_DOCX)

    doc.save(path)


def _docx_stage_script(data: dict, path: str) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _SCRIPT_FONT_DOCX
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    _docx_title_page(doc, data["project"]["title"], _SCRIPT_FONT_DOCX)

    if not data["scenes"]:
        doc.add_paragraph("No scenes.")
        doc.save(path)
        return

    current_act = None
    scene_num = 0
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if act and act != current_act:
            current_act = act
            scene_num = 0
            act_para = doc.add_paragraph()
            act_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            act_para.paragraph_format.space_before = Pt(36)
            act_run = act_para.add_run(act.upper())
            act_run.bold = True
            act_run.font.name = _SCRIPT_FONT_DOCX

        scene_num += 1
        sc_para = doc.add_paragraph()
        sc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sc_para.paragraph_format.space_before = Pt(24)
        sc_run = sc_para.add_run(f"SCENE {scene_num}")
        sc_run.bold = True
        sc_run.font.name = _SCRIPT_FONT_DOCX

        body = _scene_body(scene)
        if body:
            _add_content_paragraphs(doc, body, _SCRIPT_FONT_DOCX)

    doc.save(path)


def _docx_graphic_novel(data: dict, path: str) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _SCRIPT_FONT_DOCX
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    _docx_title_page(doc, data["project"]["title"], _SCRIPT_FONT_DOCX)

    if not data["scenes"]:
        doc.add_paragraph("No scenes.")
        doc.save(path)
        return

    page_num = 0
    for scene in data["scenes"]:
        page_num += 1
        pg_para = doc.add_paragraph()
        pg_para.paragraph_format.space_before = Pt(24)
        pg_run = pg_para.add_run(f"PAGE {page_num}")
        pg_run.bold = True
        pg_run.font.name = _SCRIPT_FONT_DOCX

        body = _scene_body(scene)
        if body:
            _add_content_paragraphs(doc, body, _SCRIPT_FONT_DOCX)

    doc.save(path)


def _add_content_paragraphs(doc, text: str, font_name: str) -> None:
    from docx.shared import Pt

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        para = doc.add_paragraph()
        lines = block.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                para.add_run().add_break()
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(12)


# -- Fountain export ----------------------------------------------------------

def build_screenplay_blocks(db: Database, project_id: int):
    """Flatten a screenplay project into one ordered ScreenplayBlock list.

    Injects a scene-heading block from the scene slug when the scene content does
    not already open with one (avoids duplicate headings on export). Read-only.
    """
    from logosforge.screenplay_blocks import ScreenplayBlock, parse_screenplay_text

    data = _gather_project_data(db, project_id)
    blocks = []
    order = 0
    for scene in data["scenes"]:
        raw = _scene_body(scene)
        parsed = parse_screenplay_text(raw, scene_id=None) if raw.strip() else []
        if not (parsed and parsed[0].element_type == "scene_heading"):
            blocks.append(ScreenplayBlock(
                element_type="scene_heading", text=_slug_line(scene),
                order_index=order))
            order += 1
        for b in parsed:
            blocks.append(ScreenplayBlock(
                element_type=b.element_type, text=b.text, order_index=order,
                metadata=dict(b.metadata)))
            order += 1
    return blocks


def export_screenplay_fountain_result(db: Database, project_id: int, *, options=None):
    """Canonical .fountain export -> FountainExportResult (Phase 10G)."""
    from logosforge.screenplay_fountain import (
        FountainExportOptions, serialize_screenplay_to_fountain,
    )
    from logosforge.screenplay_render import get_title_page, get_export_prefs

    project = db.get_project_by_id(project_id)
    title = project.title if project else "Untitled"
    if options is None:
        prefs = get_export_prefs(db, project_id)
        options = FountainExportOptions(
            include_notes=bool(prefs.get("show_notes_in_export", False)),
            include_title_page=bool(prefs.get("include_title_page", True)),
            uppercase_scene_headings=bool(prefs.get("uppercase_scene_headings", True)),
            uppercase_character_cues=bool(prefs.get("uppercase_character_cues", True)),
        )
    return serialize_screenplay_to_fountain(
        build_screenplay_blocks(db, project_id),
        title_page=get_title_page(db, project_id), options=options,
        project_title=title)


def export_screenplay_fountain(db: Database, project_id: int, *, options=None) -> str:
    """Canonical .fountain export text (Phase 10G)."""
    return export_screenplay_fountain_result(db, project_id, options=options).text


def _screenplay_render_doc(db: Database, project_id: int):
    """Render document for the professional output layer.

    Built with notes *included* so each output target decides note handling
    itself (DOCX via style.include_notes; FDX always omits + warns) rather than
    having them stripped upstream.
    """
    from logosforge.screenplay_render import build_render_document, get_export_prefs
    prefs = dict(get_export_prefs(db, project_id))
    prefs["show_notes_in_export"] = True
    return build_render_document(db, project_id, prefs=prefs)


def export_screenplay_docx(db: Database, project_id: int, path: str, *,
                           style=None, options=None):
    """Professional DOCX screenplay export (Phase 10H) -> ScreenplayDocxExportResult."""
    from logosforge.screenplay_docx_export import export_screenplay_to_docx
    from logosforge.screenplay_output_styles import get_style
    style = style or get_style()
    return export_screenplay_to_docx(
        _screenplay_render_doc(db, project_id), path, style=style, options=options)


def export_screenplay_fdx_experimental(db: Database, project_id: int, *, options=None):
    """Experimental FDX screenplay export (Phase 10H) -> ScreenplayFdxExportResult."""
    from logosforge.screenplay_fdx_export import export_screenplay_to_fdx
    return export_screenplay_to_fdx(_screenplay_render_doc(db, project_id),
                                    options=options)


def export_professional_preview_html(db: Database, project_id: int, *,
                                     dark: bool = False) -> str:
    """Professional screenplay HTML print preview (Phase 10H)."""
    from logosforge.screenplay_html_preview import build_screenplay_preview_html
    return build_screenplay_preview_html(_screenplay_render_doc(db, project_id),
                                         dark=dark)


def export_screenplay_pdf(db: Database, project_id: int, path: str, *, style=None) -> dict:
    """Approximate screenplay PDF via reportlab from the render model (Phase 10H).

    Pagination is approximate (not page-accurate). Returns a small status dict.
    """
    from logosforge.screenplay_output_styles import get_style
    style = style or get_style()
    doc = _screenplay_render_doc(db, project_id)
    warnings = ["PDF pagination is approximate (not page-accurate)."]
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except Exception as exc:
        return {"ok": False, "warnings": [f"reportlab unavailable: {exc}"]}

    def _esc(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    base = getSampleStyleSheet()["Normal"]
    base.fontName = "Courier"
    base.fontSize = style.font_size
    base.leading = style.font_size * 1.2

    def mk(name, *, left=0.0, right=0.0, bold=False, align=0, upper=False):
        return ParagraphStyle(name, parent=base, leftIndent=left * inch,
                              rightIndent=right * inch, alignment=align,
                              fontName="Courier-Bold" if bold else "Courier")

    styles = {
        "scene_heading": mk("sh", bold=True),
        "action": mk("ac"),
        "character": mk("ch", left=2.2),
        "parenthetical": mk("pa", left=1.6, right=2.0),
        "dialogue": mk("di", left=1.0, right=1.5),
        "transition": mk("tr", align=2),
        "shot": mk("st", bold=True),
        "note": mk("no"),
    }
    flow = []
    title = doc.title or (doc.title_page or {}).get("title", "")
    if title:
        flow.append(Paragraph(_esc(title.upper()), mk("title", align=1, bold=True)))
        flow.append(Spacer(1, 24))
    for b in doc.blocks:
        text = (b.export_text or b.text or "").strip()
        if not text:
            continue
        st = styles.get(b.element_type, styles["action"])
        if b.element_type in ("scene_heading", "character", "transition", "shot"):
            text = text.upper()
        flow.append(Paragraph(_esc(text), st))
    try:
        SimpleDocTemplate(path, pagesize=letter,
                          topMargin=inch, bottomMargin=inch,
                          leftMargin=1.5 * inch, rightMargin=inch).build(flow)
    except Exception as exc:
        return {"ok": False, "warnings": [f"PDF build failed: {exc}"]}
    return {"ok": True, "path": path, "warnings": warnings,
            "estimated_pages": doc.estimated_pages}


def export_screenplay_output_validation_json(db: Database, project_id: int, *,
                                             target_format: str = "docx") -> str:
    """Professional output readiness validation as JSON (Phase 10H)."""
    import datetime as _dt
    from logosforge.screenplay_output_validation import validate_professional_output
    from logosforge.writing_modes import get_project_writing_mode

    project = db.get_project_by_id(project_id)
    rep = validate_professional_output(db, project_id, target_format=target_format)
    payload = rep.to_dict()
    payload["schema_version"] = 1
    payload["writing_mode"] = get_project_writing_mode(project)
    payload["exported_at"] = _dt.datetime.now().isoformat(timespec="seconds")
    payload["project_title"] = project.title if project else "Untitled"
    return json.dumps(payload, indent=2, ensure_ascii=False)


def export_production_fountain(db: Database, project_id: int, *,
                               include_omitted: bool = True) -> str:
    """Phase 10J — Fountain with production scene numbers (#N#) + OMITTED markers.

    Opt-in; only meaningful when an active production draft with scene numbering
    exists. Default Fountain export is unchanged. Read-only, deterministic.
    """
    from logosforge.screenplay_blocks import parse_screenplay_text
    from logosforge.screenplay_fountain import (
        FountainExportOptions, serialize_screenplay_to_fountain,
    )
    from logosforge.screenplay_blocks import ScreenplayBlock
    from logosforge.screenplay_render import get_title_page
    from logosforge.screenplay_production import scene_number_map

    numbers = scene_number_map(db, project_id)
    data = _gather_project_data(db, project_id)
    blocks: list = []
    order = 0
    scenes = db.get_all_scenes(project_id)
    for scene in scenes:
        info = numbers.get(scene.id, {})
        num = info.get("number", "")
        raw = scene.content or ""
        parsed = parse_screenplay_text(raw) if raw.strip() else []
        # Heading (from content or slug) with #N# suffix.
        if parsed and parsed[0].element_type == "scene_heading":
            head = parsed[0].text
            rest = parsed[1:]
        else:
            head = (scene.slugline or scene.title or "").strip()
            rest = parsed
        if info.get("omitted"):
            if include_omitted:
                label = info.get("label") or "OMITTED"
                blocks.append(ScreenplayBlock("scene_heading",
                              f"{num} {label}".strip(), order_index=order))
                order += 1
            continue
        if head:
            head_text = f"{head} #{num}#" if num else head
            blocks.append(ScreenplayBlock("scene_heading", head_text, order_index=order))
            order += 1
        for b in rest:
            blocks.append(ScreenplayBlock(b.element_type, b.text, order_index=order))
            order += 1
    res = serialize_screenplay_to_fountain(
        blocks, title_page=get_title_page(db, project_id),
        options=FountainExportOptions(), project_title=data["project"]["title"])
    return res.text


def export_fountain_validation_json(db: Database, project_id: int) -> str:
    """Phase 10G — Fountain export validation report as JSON (read-only)."""
    import datetime as _dt
    from logosforge.screenplay_fountain import validate_fountain_export
    from logosforge.writing_modes import get_project_writing_mode

    project = db.get_project_by_id(project_id)
    res = export_screenplay_fountain_result(db, project_id)
    rep = validate_fountain_export(res.text)
    payload = rep.to_dict()
    payload["schema_version"] = 1
    payload["writing_mode"] = get_project_writing_mode(project)
    payload["exported_at"] = _dt.datetime.now().isoformat(timespec="seconds")
    payload["project_title"] = project.title if project else "Untitled"
    payload["export_warnings"] = list(res.warnings)
    payload["filename"] = res.filename
    return json.dumps(payload, indent=2, ensure_ascii=False)


def export_fountain(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)
    # Phase 10G — screenplay projects use the dedicated, canonical Fountain
    # serializer. Other modes keep the legacy multi-mode text path below.
    if _get_fmt(data) == "screenplay":
        return export_screenplay_fountain(db, project_id)

    lines: list[str] = []

    # Phase 10F — title page from project settings (falls back to project title).
    show_notes = False
    try:
        from logosforge.screenplay_render import (
            get_title_page, get_export_prefs, title_page_to_fountain,
        )
        prefs = get_export_prefs(db, project_id)
        show_notes = bool(prefs.get("show_notes_in_export", False))
        tp_lines = []
        if prefs.get("include_title_page", True):
            tp_lines = title_page_to_fountain(get_title_page(db, project_id))
        if tp_lines:
            lines.extend(tp_lines)
        else:
            lines.append(f"Title: {data['project']['title']}")
    except Exception:
        lines.append(f"Title: {data['project']['title']}")
    lines.append("")
    lines.append("")

    fmt = _get_fmt(data)

    current_act = None
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if act and act != current_act:
            current_act = act
            if fmt in ("series", "stage_script"):
                lines.append(f"= {act}")
                lines.append("")

        # Avoid a duplicate heading: only emit the metadata slug when the scene
        # content doesn't already start with its own scene heading.
        if not _scene_starts_with_heading(scene):
            slug = _slug_line(scene)
            lines.append(f".{slug}")
            lines.append("")

        body = _screenplay_body(scene, fountain=True, include_notes=show_notes)
        if body:
            lines.append(body)
            lines.append("")

    return "\n".join(lines)


def export_screenplay_preview_html(db: Database, project_id: int) -> str:
    """Phase 10F — conservative screenplay preview HTML (not page-accurate)."""
    from logosforge.screenplay_render import build_render_document, render_to_html
    return render_to_html(build_render_document(db, project_id))


def export_screenplay_export_validation_json(db: Database, project_id: int,
                                             *, target_format: str = "fountain") -> str:
    """Phase 10F — export-readiness validation as JSON (read-only)."""
    import datetime as _dt
    from logosforge.screenplay_export_validation import validate_screenplay_export
    from logosforge.screenplay_render import get_export_prefs
    from logosforge.writing_modes import get_project_writing_mode

    project = db.get_project_by_id(project_id)
    report = validate_screenplay_export(
        db, project_id, target_format=target_format,
        prefs=get_export_prefs(db, project_id))
    payload = report.to_dict()
    payload["schema_version"] = 1
    payload["writing_mode"] = get_project_writing_mode(project)
    payload["exported_at"] = _dt.datetime.now().isoformat(timespec="seconds")
    payload["project_title"] = project.title if project else "Untitled"
    return json.dumps(payload, indent=2, ensure_ascii=False)


# -- FDX (Final Draft XML) export ---------------------------------------------

def export_fdx(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)

    root = ET.Element("FinalDraft", DocumentType="Script", Template="No", Version="4")
    content = ET.SubElement(root, "Content")

    title_para = ET.SubElement(content, "Paragraph", Type="Action")
    title_text = ET.SubElement(title_para, "Text")
    title_text.text = data["project"]["title"]

    fmt = _get_fmt(data)

    current_act = None
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if act and act != current_act:
            current_act = act
            if fmt in ("series", "stage_script"):
                act_para = ET.SubElement(content, "Paragraph", Type="Action")
                act_t = ET.SubElement(act_para, "Text")
                act_t.text = act.upper()

        slug = _slug_line(scene)
        slug_para = ET.SubElement(content, "Paragraph", Type="Scene Heading")
        slug_text = ET.SubElement(slug_para, "Text")
        slug_text.text = slug

        body = _scene_body(scene)
        if body:
            for block in body.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                action_para = ET.SubElement(content, "Paragraph", Type="Action")
                action_text = ET.SubElement(action_para, "Text")
                action_text.text = block.replace("\n", " ")

    tree = ET.ElementTree(root)
    buf = io.BytesIO()
    tree.write(buf, encoding="utf-8", xml_declaration=True)
    return buf.getvalue().decode("utf-8")


# -- PDF export ----------------------------------------------------------------

def export_pdf(db: Database, project_id: int, path: str) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
    )

    data = _gather_project_data(db, project_id)
    fmt = _get_fmt(data)
    is_script = _is_script_format(fmt)
    font_name = "Courier" if is_script else "Times-Roman"
    font_bold = "Courier-Bold" if is_script else "Times-Bold"
    font_italic = "Courier-Oblique" if is_script else "Times-Italic"

    doc = SimpleDocTemplate(
        path,
        pagesize=letter,
        leftMargin=1.5 * inch if is_script else 1.0 * inch,
        rightMargin=1.0 * inch,
        topMargin=1.0 * inch,
        bottomMargin=1.0 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ScriptTitle",
        parent=styles["Title"],
        fontName=font_bold,
        fontSize=24,
        alignment=1,
        spaceAfter=36,
    )
    heading_style = ParagraphStyle(
        "SceneHeading",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=12,
        spaceBefore=24,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=12,
        leading=14 if is_script else 16,
        spaceAfter=6 if is_script else 8,
    )
    act_style = ParagraphStyle(
        "Act",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=14,
        alignment=1,
        spaceBefore=24,
        spaceAfter=12,
    )
    chapter_style = ParagraphStyle(
        "Chapter",
        parent=styles["Normal"],
        fontName=font_bold,
        fontSize=16,
        spaceBefore=24,
        spaceAfter=12,
    )

    elements: list = []
    elements.append(Paragraph(_esc(data["project"]["title"]), title_style))
    elements.append(Spacer(1, 36))

    if not data["scenes"]:
        elements.append(Paragraph("No scenes.", body_style))
        doc.build(elements)
        return

    if fmt in ("screenplay", "series"):
        _pdf_screenplay(data, fmt, elements, heading_style, body_style, act_style)
    elif fmt == "stage_script":
        _pdf_stage_script(data, elements, heading_style, body_style, act_style)
    elif fmt == "graphic_novel":
        _pdf_graphic_novel(data, elements, heading_style, body_style)
    else:
        _pdf_novel(data, elements, chapter_style, heading_style, body_style)

    doc.build(elements)


def _esc(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _pdf_body_blocks(text: str, style, elements: list) -> None:
    from reportlab.platypus import Paragraph

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        safe = _esc(block).replace("\n", "<br/>")
        elements.append(Paragraph(safe, style))


def _pdf_novel(data, elements, chapter_style, heading_style, body_style):
    from reportlab.platypus import Paragraph

    chapter_groups = _group_scenes_by_chapter(data["scenes"])
    chapter_num = 0
    for chapter_name, group_scenes in chapter_groups:
        if chapter_name:
            chapter_num += 1
            elements.append(
                Paragraph(_esc(f"Chapter {chapter_num}: {chapter_name}"), chapter_style),
            )
        for scene in group_scenes:
            elements.append(Paragraph(f"<i>{_esc(scene['title'])}</i>", heading_style))
            body = _scene_body(scene)
            if body:
                _pdf_body_blocks(body, body_style, elements)


def _pdf_screenplay(data, fmt, elements, heading_style, body_style, act_style):
    from reportlab.platypus import Paragraph

    current_act = None
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if fmt == "series" and act and act != current_act:
            current_act = act
            elements.append(Paragraph(_esc(act.upper()), act_style))
        elements.append(Paragraph(_esc(_slug_line(scene)), heading_style))
        body = _scene_body(scene)
        if body:
            _pdf_body_blocks(body, body_style, elements)


def _pdf_stage_script(data, elements, heading_style, body_style, act_style):
    from reportlab.platypus import Paragraph

    current_act = None
    scene_num = 0
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if act and act != current_act:
            current_act = act
            scene_num = 0
            elements.append(Paragraph(_esc(act.upper()), act_style))
        scene_num += 1
        elements.append(Paragraph(_esc(f"SCENE {scene_num}"), heading_style))
        body = _scene_body(scene)
        if body:
            _pdf_body_blocks(body, body_style, elements)


def _pdf_graphic_novel(data, elements, heading_style, body_style):
    from reportlab.platypus import Paragraph

    page_num = 0
    for scene in data["scenes"]:
        page_num += 1
        elements.append(Paragraph(_esc(f"PAGE {page_num}"), heading_style))
        body = _scene_body(scene)
        if body:
            _pdf_body_blocks(body, body_style, elements)


# -- HTML export ---------------------------------------------------------------

def export_html(db: Database, project_id: int) -> str:
    data = _gather_project_data(db, project_id)
    fmt = _get_fmt(data)
    is_script = _is_script_format(fmt)
    font = "Courier New, Courier, monospace" if is_script else "Times New Roman, Georgia, serif"
    title = _esc(data["project"]["title"])

    css = (
        "body { max-width: 720px; margin: 40px auto; padding: 0 20px; "
        f"font-family: {font}; font-size: 12pt; line-height: 1.5; }}\n"
        "h1 { text-align: center; }\n"
        "h2 { text-align: center; text-transform: uppercase; font-size: 14pt; }\n"
        ".scene-heading { font-weight: bold; text-transform: uppercase; "
        "margin-top: 24px; margin-bottom: 12px; }\n"
        ".chapter { font-size: 16pt; font-weight: bold; margin-top: 36px; }\n"
        ".body-text { margin-bottom: 8px; }\n"
    )

    parts: list[str] = [
        "<!DOCTYPE html>",
        "<html lang=\"en\">",
        "<head>",
        f"<meta charset=\"utf-8\"><title>{title}</title>",
        f"<style>{css}</style>",
        "</head>",
        "<body>",
        f"<h1>{title}</h1>",
    ]

    if not data["scenes"]:
        parts.append("<p>No scenes.</p>")
    elif fmt in ("screenplay", "series"):
        _html_screenplay(data, fmt, parts)
    elif fmt == "stage_script":
        _html_stage_script(data, parts)
    elif fmt == "graphic_novel":
        _html_graphic_novel(data, parts)
    else:
        _html_novel(data, parts)

    parts.append("</body>")
    parts.append("</html>")
    return "\n".join(parts)


def _html_body_blocks(text: str, parts: list[str]) -> None:
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        safe = _esc(block).replace("\n", "<br>")
        parts.append(f"<p class=\"body-text\">{safe}</p>")


def _html_novel(data, parts):
    chapter_groups = _group_scenes_by_chapter(data["scenes"])
    chapter_num = 0
    for chapter_name, group_scenes in chapter_groups:
        if chapter_name:
            chapter_num += 1
            parts.append(f"<div class=\"chapter\">Chapter {chapter_num}: {_esc(chapter_name)}</div>")
        for scene in group_scenes:
            parts.append(f"<div class=\"scene-heading\"><em>{_esc(scene['title'])}</em></div>")
            body = _scene_body(scene)
            if body:
                _html_body_blocks(body, parts)


def _html_screenplay(data, fmt, parts):
    current_act = None
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if fmt == "series" and act and act != current_act:
            current_act = act
            parts.append(f"<h2>{_esc(act.upper())}</h2>")
        parts.append(f"<div class=\"scene-heading\">{_esc(_slug_line(scene))}</div>")
        body = _scene_body(scene)
        if body:
            _html_body_blocks(body, parts)


def _html_stage_script(data, parts):
    current_act = None
    scene_num = 0
    for scene in data["scenes"]:
        act = scene.get("act", "")
        if act and act != current_act:
            current_act = act
            scene_num = 0
            parts.append(f"<h2>{_esc(act.upper())}</h2>")
        scene_num += 1
        parts.append(f"<div class=\"scene-heading\">SCENE {scene_num}</div>")
        body = _scene_body(scene)
        if body:
            _html_body_blocks(body, parts)


def _html_graphic_novel(data, parts):
    page_num = 0
    for scene in data["scenes"]:
        page_num += 1
        parts.append(f"<div class=\"scene-heading\">PAGE {page_num}</div>")
        body = _scene_body(scene)
        if body:
            _html_body_blocks(body, parts)
