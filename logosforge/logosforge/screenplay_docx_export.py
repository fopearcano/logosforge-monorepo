"""Screenplay DOCX export (Phase 10H).

Builds a clean, readable, structurally-faithful screenplay ``.docx`` from a
:class:`~logosforge.screenplay_render.ScreenplayRenderDocument` using a
:class:`~logosforge.screenplay_output_styles.ScreenplayOutputStyle`. Not
Final Draft-perfect; deterministic; no DB mutation, no LLM. Degrades cleanly if
``python-docx`` is unavailable.
"""

from __future__ import annotations

import datetime as _dt
from dataclasses import dataclass, field
from typing import Any

from logosforge.screenplay_output_styles import (
    DEFAULT_STYLE, ScreenplayOutputStyle,
)


def docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


@dataclass
class ScreenplayDocxExportResult:
    file_path: str | None = None
    ok: bool = False
    warnings: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    estimated_pages: float | None = None
    exported_at: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "file_path": self.file_path, "ok": self.ok,
            "warnings": list(self.warnings), "metadata": dict(self.metadata),
            "estimated_pages": self.estimated_pages, "exported_at": self.exported_at,
        }


def export_screenplay_to_docx(
    render_document, path: str, *,
    style: ScreenplayOutputStyle | None = None, options: dict | None = None,
) -> ScreenplayDocxExportResult:
    """Render a screenplay document to a ``.docx`` file at *path*."""
    style = style or DEFAULT_STYLE
    options = options or {}
    result = ScreenplayDocxExportResult(
        file_path=path, exported_at=_dt.datetime.now().isoformat(timespec="seconds"))

    if not docx_available():
        result.warnings.append(
            "python-docx is not installed — DOCX export unavailable.")
        return result

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _ALIGN = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }

    doc = Document()
    # Base font + margins.
    normal = doc.styles["Normal"]
    normal.font.name = style.font_family
    normal.font.size = Pt(style.font_size)
    for section in doc.sections:
        section.top_margin = Inches(style.margins_in["top"])
        section.bottom_margin = Inches(style.margins_in["bottom"])
        section.left_margin = Inches(style.margins_in["left"])
        section.right_margin = Inches(style.margins_in["right"])

    # -- Title page --
    tp = getattr(render_document, "title_page", {}) or {}
    if style.include_title_page and (tp.get("title") or render_document.title):
        for key in ("title", "credit", "author", "source", "draft_date", "contact"):
            val = tp.get(key) or (render_document.title if key == "title" else "")
            if val:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(val).upper() if key == "title" else str(val))
                run.bold = (key == "title")
        doc.add_page_break()
    elif style.include_title_page:
        result.warnings.append("No title page metadata — exported without one.")

    def _add(text: str, est: "ElementStyle", note: bool = False) -> None:
        p = doc.add_paragraph()
        p.alignment = _ALIGN.get(est.align, WD_ALIGN_PARAGRAPH.LEFT)
        pf = p.paragraph_format
        if est.left_indent_in:
            pf.left_indent = Inches(est.left_indent_in)
        if est.right_indent_in:
            pf.right_indent = Inches(est.right_indent_in)
        if est.space_before_pt:
            pf.space_before = Pt(est.space_before_pt)
        run = p.add_run(text.upper() if est.uppercase else text)
        run.bold = est.bold
        run.italic = est.italic
        if note:
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    notes_omitted = 0
    block_count = 0
    for b in render_document.blocks:
        et = b.element_type
        text = (b.export_text or b.text or "").strip()
        if not text:
            continue
        if et == "note" and not style.include_notes:
            notes_omitted += 1
            continue
        _add(text, style.style_for(et), note=(et == "note"))
        block_count += 1

    if notes_omitted:
        result.warnings.append(
            f"{notes_omitted} note block(s) omitted (style.include_notes=off).")
    result.warnings.extend(render_document.warnings or [])

    try:
        doc.save(path)
        result.ok = True
    except Exception as exc:
        result.warnings.append(f"Failed to write DOCX: {exc}")
        return result

    result.estimated_pages = getattr(render_document, "estimated_pages", None)
    result.metadata = {"block_count": block_count, "font": style.font_family,
                       "notes_omitted": notes_omitted}
    return result
