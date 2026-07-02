"""Tests for format-aware export and import."""

import json
import xml.etree.ElementTree as ET

from logosforge.db import Database
from logosforge.export import (
    export_csv_scenes,
    export_docx_manuscript,
    export_fdx,
    export_formatted_text,
    export_fountain,
    export_html,
    export_json,
    export_manuscript,
    export_markdown,
    export_pdf,
    export_screenplay,
)
from logosforge.import_data import import_json, validate_import_data


def _make_project(db, fmt="novel"):
    proj = db.create_project("Test Story", format_mode=fmt)
    db.create_scene(proj.id, "Opening", content="The dawn broke.", act="Act One", chapter="Chapter 1")
    db.create_scene(proj.id, "Climax", content="She faced her fears.", act="Act Two", chapter="Chapter 2")
    return proj


# -- JSON round-trip preserves format_mode ------------------------------------

def test_json_export_includes_format_mode():
    db = Database()
    proj = _make_project(db, "screenplay")
    raw = export_json(db, proj.id)
    data = json.loads(raw)
    assert data["project"]["format_mode"] == "screenplay"


def test_json_export_novel_format_mode():
    db = Database()
    proj = _make_project(db, "novel")
    raw = export_json(db, proj.id)
    data = json.loads(raw)
    assert data["project"]["format_mode"] == "novel"


def test_json_roundtrip_format_mode():
    db = Database()
    proj = _make_project(db, "stage_script")
    raw = export_json(db, proj.id)
    data, err = validate_import_data(raw)
    assert err == ""
    new_id = import_json(db, data)
    new_proj = db.get_project_by_id(new_id)
    assert new_proj.format_mode == "stage_script"


def test_json_roundtrip_graphic_novel():
    db = Database()
    proj = _make_project(db, "graphic_novel")
    raw = export_json(db, proj.id)
    new_id = import_json(db, json.loads(raw))
    new_proj = db.get_project_by_id(new_id)
    assert new_proj.format_mode == "graphic_novel"


def test_json_roundtrip_series():
    db = Database()
    proj = _make_project(db, "series")
    raw = export_json(db, proj.id)
    new_id = import_json(db, json.loads(raw))
    new_proj = db.get_project_by_id(new_id)
    assert new_proj.format_mode == "series"


def test_import_missing_format_mode_defaults_novel():
    db = Database()
    data = {
        "project": {"title": "Legacy"},
        "characters": [],
        "places": [],
        "notes": [],
        "scenes": [],
    }
    new_id = import_json(db, data)
    proj = db.get_project_by_id(new_id)
    assert proj.format_mode == "novel"


# -- Format-aware text export ------------------------------------------------

def test_formatted_text_novel():
    db = Database()
    proj = _make_project(db, "novel")
    text = export_formatted_text(db, proj.id)
    assert "Test Story" in text
    assert "The dawn broke." in text
    assert "Chapter 1" in text


def test_formatted_text_screenplay():
    db = Database()
    proj = _make_project(db, "screenplay")
    text = export_formatted_text(db, proj.id)
    assert "TEST STORY" in text
    assert "OPENING" in text
    assert "The dawn broke." in text


def test_formatted_text_stage_script():
    db = Database()
    proj = _make_project(db, "stage_script")
    text = export_formatted_text(db, proj.id)
    assert "TEST STORY" in text
    assert "ACT ONE" in text
    assert "SCENE" in text


def test_formatted_text_graphic_novel():
    db = Database()
    proj = _make_project(db, "graphic_novel")
    text = export_formatted_text(db, proj.id)
    assert "PAGE 1" in text
    assert "PAGE 2" in text
    assert "The dawn broke." in text


def test_formatted_text_series():
    db = Database()
    proj = _make_project(db, "series")
    text = export_formatted_text(db, proj.id)
    assert "TEST STORY" in text
    assert "ACT ONE" in text
    assert "ACT TWO" in text


# -- Legacy exports still work -----------------------------------------------

def test_export_screenplay_function():
    db = Database()
    proj = _make_project(db)
    text = export_screenplay(db, proj.id)
    assert "TEST STORY" in text


def test_export_manuscript_function():
    db = Database()
    proj = _make_project(db)
    text = export_manuscript(db, proj.id)
    assert "Test Story" in text
    assert "Chapter 1" in text


# -- DOCX format-aware -------------------------------------------------------

def test_docx_novel(tmp_path):
    db = Database()
    proj = _make_project(db, "novel")
    path = str(tmp_path / "novel.docx")
    export_docx_manuscript(db, proj.id, path)
    from docx import Document
    doc = Document(path)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Test Story" in full
    assert "The dawn broke." in full


def test_docx_screenplay(tmp_path):
    db = Database()
    proj = _make_project(db, "screenplay")
    path = str(tmp_path / "screenplay.docx")
    export_docx_manuscript(db, proj.id, path)
    from docx import Document
    doc = Document(path)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Test Story" in full
    assert "OPENING" in full


def test_docx_stage_script(tmp_path):
    db = Database()
    proj = _make_project(db, "stage_script")
    path = str(tmp_path / "stage.docx")
    export_docx_manuscript(db, proj.id, path)
    from docx import Document
    doc = Document(path)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Test Story" in full
    assert "ACT ONE" in full
    assert "SCENE 1" in full


def test_docx_graphic_novel(tmp_path):
    db = Database()
    proj = _make_project(db, "graphic_novel")
    path = str(tmp_path / "graphic.docx")
    export_docx_manuscript(db, proj.id, path)
    from docx import Document
    doc = Document(path)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "PAGE 1" in full


def test_docx_series(tmp_path):
    db = Database()
    proj = _make_project(db, "series")
    path = str(tmp_path / "series.docx")
    export_docx_manuscript(db, proj.id, path)
    from docx import Document
    doc = Document(path)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Test Story" in full
    assert "ACT ONE" in full


# -- Other exports unaffected ------------------------------------------------

def test_markdown_export():
    db = Database()
    proj = _make_project(db)
    text = export_markdown(db, proj.id)
    assert "# Test Story" in text
    assert "## Scenes" in text


def test_csv_export():
    db = Database()
    proj = _make_project(db)
    text = export_csv_scenes(db, proj.id)
    assert "Opening" in text
    assert "Climax" in text


def test_json_roundtrip_scenes_preserved():
    db = Database()
    proj = _make_project(db, "screenplay")
    raw = export_json(db, proj.id)
    data = json.loads(raw)
    assert len(data["scenes"]) == 2
    new_id = import_json(db, data)
    scenes = db.get_all_scenes(new_id)
    assert len(scenes) == 2
    assert scenes[0].content == "The dawn broke."
    assert scenes[1].content == "She faced her fears."


# -- Fountain export ----------------------------------------------------------

def test_fountain_screenplay():
    db = Database()
    proj = _make_project(db, "screenplay")
    text = export_fountain(db, proj.id)
    assert "Title: Test Story" in text
    assert ".OPENING" in text
    assert "The dawn broke." in text


def test_fountain_novel():
    db = Database()
    proj = _make_project(db, "novel")
    text = export_fountain(db, proj.id)
    assert "Title: Test Story" in text
    assert ".OPENING" in text


def test_fountain_series_has_acts():
    db = Database()
    proj = _make_project(db, "series")
    text = export_fountain(db, proj.id)
    assert "= Act One" in text
    assert "= Act Two" in text


def test_fountain_stage_script_has_acts():
    db = Database()
    proj = _make_project(db, "stage_script")
    text = export_fountain(db, proj.id)
    assert "= Act One" in text


# -- FDX (Final Draft XML) export ---------------------------------------------

def test_fdx_valid_xml():
    db = Database()
    proj = _make_project(db, "screenplay")
    xml_str = export_fdx(db, proj.id)
    root = ET.fromstring(xml_str)
    assert root.tag == "FinalDraft"
    assert root.attrib["DocumentType"] == "Script"


def test_fdx_contains_scenes():
    db = Database()
    proj = _make_project(db, "screenplay")
    xml_str = export_fdx(db, proj.id)
    root = ET.fromstring(xml_str)
    headings = [
        p for p in root.iter("Paragraph")
        if p.attrib.get("Type") == "Scene Heading"
    ]
    assert len(headings) == 2
    assert "OPENING" in headings[0].find("Text").text


def test_fdx_series_includes_acts():
    db = Database()
    proj = _make_project(db, "series")
    xml_str = export_fdx(db, proj.id)
    root = ET.fromstring(xml_str)
    actions = [
        p.find("Text").text
        for p in root.iter("Paragraph")
        if p.attrib.get("Type") == "Action" and p.find("Text").text
    ]
    assert any("ACT ONE" in a for a in actions)


# -- PDF export ----------------------------------------------------------------

def test_pdf_novel(tmp_path):
    db = Database()
    proj = _make_project(db, "novel")
    path = str(tmp_path / "novel.pdf")
    export_pdf(db, proj.id, path)
    with open(path, "rb") as f:
        header = f.read(5)
    assert header == b"%PDF-"


def test_pdf_screenplay(tmp_path):
    db = Database()
    proj = _make_project(db, "screenplay")
    path = str(tmp_path / "screenplay.pdf")
    export_pdf(db, proj.id, path)
    with open(path, "rb") as f:
        header = f.read(5)
    assert header == b"%PDF-"


def test_pdf_stage_script(tmp_path):
    db = Database()
    proj = _make_project(db, "stage_script")
    path = str(tmp_path / "stage.pdf")
    export_pdf(db, proj.id, path)
    with open(path, "rb") as f:
        header = f.read(5)
    assert header == b"%PDF-"


def test_pdf_graphic_novel(tmp_path):
    db = Database()
    proj = _make_project(db, "graphic_novel")
    path = str(tmp_path / "graphic.pdf")
    export_pdf(db, proj.id, path)
    with open(path, "rb") as f:
        header = f.read(5)
    assert header == b"%PDF-"


def test_pdf_series(tmp_path):
    db = Database()
    proj = _make_project(db, "series")
    path = str(tmp_path / "series.pdf")
    export_pdf(db, proj.id, path)
    with open(path, "rb") as f:
        header = f.read(5)
    assert header == b"%PDF-"


def test_pdf_empty_project(tmp_path):
    db = Database()
    proj = db.create_project("Empty", format_mode="novel")
    path = str(tmp_path / "empty.pdf")
    export_pdf(db, proj.id, path)
    with open(path, "rb") as f:
        header = f.read(5)
    assert header == b"%PDF-"


# -- HTML export ---------------------------------------------------------------

def test_html_novel():
    db = Database()
    proj = _make_project(db, "novel")
    html = export_html(db, proj.id)
    assert "<!DOCTYPE html>" in html
    assert "Test Story" in html
    assert "The dawn broke." in html
    assert "Chapter 1" in html


def test_html_screenplay():
    db = Database()
    proj = _make_project(db, "screenplay")
    html = export_html(db, proj.id)
    assert "<!DOCTYPE html>" in html
    assert "OPENING" in html
    assert "monospace" in html


def test_html_stage_script():
    db = Database()
    proj = _make_project(db, "stage_script")
    html = export_html(db, proj.id)
    assert "ACT ONE" in html
    assert "SCENE 1" in html


def test_html_graphic_novel():
    db = Database()
    proj = _make_project(db, "graphic_novel")
    html = export_html(db, proj.id)
    assert "PAGE 1" in html
    assert "PAGE 2" in html


def test_html_series():
    db = Database()
    proj = _make_project(db, "series")
    html = export_html(db, proj.id)
    assert "ACT ONE" in html
    assert "ACT TWO" in html


def test_html_empty_project():
    db = Database()
    proj = db.create_project("Empty")
    html = export_html(db, proj.id)
    assert "No scenes." in html
